import io
import re
import pdfplumber # type: ignore
from docx import Document # type: ignore
import pytesseract # type: ignore
from pdf2image import convert_from_bytes # type: ignore


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]
    try:
        if ext == "pdf":
            return _extract_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return _extract_from_docx(file_bytes)
        else:
            return _extract_from_txt(file_bytes)
    except Exception:
        return ""


def _extract_from_pdf(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text(layout=True) or "" for p in pdf.pages]
            text = "\n".join(pages).strip()
            if not text:
                text = _ocr_fallback(file_bytes)
            return text
    except Exception:
        return ""


def _ocr_fallback(file_bytes: bytes) -> str:
    try:
        images = convert_from_bytes(file_bytes)
        return "\n".join(pytesseract.image_to_string(img) for img in images).strip()
    except Exception:
        return ""


def _extract_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for r in t.rows:
                parts.extend([c.text for c in r.cells])
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _extract_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def extract_entities(text: str) -> dict:
    if not text:
        return {"email": "Not Found", "phone": "Not Found"}
    email = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    return {
        "email": email.group(0) if email else "Not Found",
        "phone": phone.group(0) if phone else "Not Found",
    }
